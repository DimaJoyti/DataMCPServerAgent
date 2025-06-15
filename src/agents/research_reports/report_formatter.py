"""
Report formatter component for the research reports agent.
This module provides functionality for formatting research reports.
"""

import os
import time
from datetime import datetime
from typing import Any, Dict

from langchain_anthropic import ChatAnthropic

from src.memory.memory_persistence import MemoryDatabase
from src.utils.error_handlers import format_error_for_user


class ReportFormatter:
    """Component for formatting research reports."""

    def __init__(self, model: ChatAnthropic, memory_db: MemoryDatabase):
        """Initialize the report formatter.

        Args:
            model: Language model to use
            memory_db: Memory database for persistence
        """
        self.model = model
        self.memory_db = memory_db

    async def format_report(
        self, report: Dict[str, Any], format_type: str = "markdown", filename: str = None
    ) -> Dict[str, Any]:
        """Format a research report.

        Args:
            report: Research report
            format_type: Output format type
            filename: Output filename

        Returns:
            Formatted report information
        """
        try:
            if filename is None:
                # Generate a filename based on the topic
                topic_slug = report["topic"].lower().replace(" ", "_").replace("/", "_")
                timestamp = int(time.time())
                filename = f"{topic_slug}_{timestamp}.{self._get_extension(format_type)}"

            # Create the reports directory if it doesn't exist
            os.makedirs("reports", exist_ok=True)
            filepath = os.path.join("reports", filename)

            # Format the report based on the format type
            if format_type == "markdown":
                content = self._format_markdown(report)
                self._save_to_file(content, filepath)
            elif format_type == "html":
                content = self._format_html(report)
                self._save_to_file(content, filepath)
            elif format_type == "pdf":
                content = self._format_pdf(report, filepath)
            elif format_type == "docx":
                content = self._format_docx(report, filepath)
            else:
                raise ValueError(f"Unsupported format type: {format_type}")

            # Store formatted report in memory
            self.memory_db.save_entity(
                "formatted_reports",
                f"{format_type}_{int(time.time())}",
                {
                    "report_id": report.get("metadata", {}).get("generated_at"),
                    "format_type": format_type,
                    "filename": filename,
                    "filepath": filepath,
                    "content": content if format_type in ["markdown", "html"] else None,
                },
            )

            return {
                "format_type": format_type,
                "filename": filename,
                "filepath": filepath,
                "content": content if format_type in ["markdown", "html"] else None,
            }
        except Exception as e:
            error_message = format_error_for_user(e)
            print(f"Error formatting report: {error_message}")
            return {"error": error_message}

    def _get_extension(self, format_type: str) -> str:
        """Get the file extension for a format type.

        Args:
            format_type: Format type

        Returns:
            File extension
        """
        extensions = {"markdown": "md", "html": "html", "pdf": "pdf", "docx": "docx"}
        return extensions.get(format_type, "txt")

    def _format_markdown(self, report: Dict[str, Any]) -> str:
        """Format a report as Markdown.

        Args:
            report: Research report

        Returns:
            Markdown content
        """
        # Create Markdown content
        markdown = f"# {report['topic']}\n\n"

        # Add executive summary
        markdown += "## Executive Summary\n\n"
        markdown += f"{report['executive_summary']}\n\n"

        # Add sections
        for section_name, section_content in report["sections"].items():
            markdown += f"## {section_name}\n\n"
            markdown += f"{section_content}\n\n"

        # Add bibliography
        markdown += "## Bibliography\n\n"
        for entry in report["bibliography"]:
            markdown += f"- {entry}\n"

        # Add timestamp
        markdown += f"\n\n*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n"

        return markdown

    def _format_html(self, report: Dict[str, Any]) -> str:
        """Format a report as HTML.

        Args:
            report: Research report

        Returns:
            HTML content
        """
        # Create HTML content
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{report['topic']}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1, h2 {{
            color: #333;
        }}
        .executive-summary {{
            background-color: #f5f5f5;
            padding: 15px;
            border-left: 4px solid #333;
            margin-bottom: 20px;
        }}
        .bibliography {{
            margin-top: 30px;
            border-top: 1px solid #ddd;
            padding-top: 20px;
        }}
        .timestamp {{
            font-style: italic;
            color: #777;
            margin-top: 30px;
            font-size: 0.9em;
        }}
    </style>
</head>
<body>
    <h1>{report['topic']}</h1>

    <div class="executive-summary">
        <h2>Executive Summary</h2>
"""
        # Add executive summary with paragraph handling
        executive_summary = report["executive_summary"]
        paragraphs = executive_summary.split("\n\n")
        for paragraph in paragraphs:
            html += f"        <p>{paragraph}</p>\n"

        html += "    </div>"

        # Add sections
        for section_name, section_content in report["sections"].items():
            html += f"""    <div class="section">
        <h2>{section_name}</h2>
"""
            # Add section content with paragraph handling
            paragraphs = section_content.split("\n\n")
            for paragraph in paragraphs:
                html += f"        <p>{paragraph}</p>\n"

            html += "    </div>\n"

        # Add bibliography
        html += """    <div class="bibliography">
        <h2>Bibliography</h2>
        <ul>
"""
        for entry in report["bibliography"]:
            html += f"            <li>{entry}</li>\n"

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        html += f"""        </ul>
    </div>

    <div class="timestamp">
        <p>Generated on {timestamp}</p>
    </div>
</body>
</html>"""

        return html

    def _format_pdf(self, report: Dict[str, Any], filepath: str) -> str:
        """Format a report as PDF.

        Args:
            report: Research report
            filepath: Output file path

        Returns:
            Success message
        """
        try:
            # First, generate HTML
            html_content = self._format_html(report)

            # Try to use a PDF library if available
            try:
                from weasyprint import HTML

                HTML(string=html_content).write_pdf(filepath)
                return f"PDF report saved to {filepath}"
            except ImportError:
                # Fall back to a simpler approach
                html_filepath = filepath.replace(".pdf", ".html")
                self._save_to_file(html_content, html_filepath)
                return f"HTML report saved to {html_filepath} (PDF conversion not available)"
        except Exception as e:
            return f"Error creating PDF: {str(e)}"

    def _format_docx(self, report: Dict[str, Any], filepath: str) -> str:
        """Format a report as DOCX.

        Args:
            report: Research report
            filepath: Output file path

        Returns:
            Success message
        """
        try:
            # Try to use python-docx if available
            try:
                from docx import Document
                from docx.shared import Inches, Pt

                document = Document()

                # Add title
                document.add_heading(report["topic"], 0)

                # Add executive summary
                document.add_heading("Executive Summary", 1)
                document.add_paragraph(report["executive_summary"])

                # Add sections
                for section_name, section_content in report["sections"].items():
                    document.add_heading(section_name, 1)
                    document.add_paragraph(section_content)

                # Add bibliography
                document.add_heading("Bibliography", 1)
                for entry in report["bibliography"]:
                    document.add_paragraph(entry, style="List Bullet")

                # Add timestamp
                document.add_paragraph(
                    f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style="Subtitle"
                )

                # Save the document
                document.save(filepath)
                return f"DOCX report saved to {filepath}"
            except ImportError:
                # Fall back to a simpler approach
                markdown_content = self._format_markdown(report)
                markdown_filepath = filepath.replace(".docx", ".md")
                self._save_to_file(markdown_content, markdown_filepath)
                return (
                    f"Markdown report saved to {markdown_filepath} (DOCX conversion not available)"
                )
        except Exception as e:
            return f"Error creating DOCX: {str(e)}"

    def _save_to_file(self, content: str, filepath: str) -> None:
        """Save content to a file.

        Args:
            content: Content to save
            filepath: Output file path
        """
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
