"""
DOCX file parser implementation.
"""

import logging
from pathlib import Path
from typing import List, Union, Dict, Any

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from .base_parser import BaseParser, ParsedDocument
from ..metadata.models import DocumentMetadata, DocumentType
from ..metadata.extractor import MetadataExtractor

class DOCXParser(BaseParser):
    """Parser for DOCX files."""

    def __init__(self, *args, **kwargs):
        """Initialize DOCX parser."""
        super().__init__(*args, **kwargs)

        if not HAS_DOCX:
            raise ImportError(
                "DOCX parsing requires python-docx. "
                "Install with: pip install python-docx"
            )

    @property
    def supported_types(self) -> List[DocumentType]:
        """Return list of supported document types."""
        return [DocumentType.DOCX]

    @property
    def supported_extensions(self) -> List[str]:
        """Return list of supported file extensions."""
        return ['docx', 'docm']

    def _parse_file_impl(self, file_path: Path) -> ParsedDocument:
        """
        Parse a DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            ParsedDocument: Parsed document result
        """
        # Extract basic metadata
        extractor = MetadataExtractor()
        metadata = extractor.extract_from_file(file_path)
        metadata.document_type = DocumentType.DOCX

        text_content = []
        tables = []
        images = []
        warnings = []
        errors = []

        try:
            # Open DOCX document
            doc = Document(file_path)

            # Extract document properties
            self._extract_docx_properties(doc, metadata)

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract tables if configured
            if self.config.extract_tables:
                for table_idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)

                    tables.append({
                        'table_index': table_idx,
                        'data': table_data,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0
                    })

            # Extract images if configured
            if self.config.extract_images:
                # Get image relationships
                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        images.append({
                            'relationship_id': rel.rId,
                            'target': rel.target_ref,
                            'type': rel.reltype
                        })

        except Exception as e:
            error_msg = f"Error parsing DOCX file: {str(e)}"
            errors.append(error_msg)
            if not self.config.ignore_errors:
                raise

        # Combine all text
        full_text = '\n\n'.join(text_content)

        # Normalize text if configured
        if self.config.normalize_whitespace:
            full_text = self._normalize_text(full_text)

        # Extract links
        links = self._extract_links(full_text)

        # Update text statistics
        if full_text:
            metadata.character_count = len(full_text)
            metadata.word_count = len(full_text.split())

            # Count paragraphs (non-empty)
            paragraphs = [p for p in full_text.split('\n\n') if p.strip()]
            metadata.paragraph_count = len(paragraphs)

        # Create result
        result = ParsedDocument(
            text=full_text,
            metadata=metadata,
            tables=tables,
            images=images,
            links=links,
            warnings=warnings,
            errors=errors,
            parsing_time=0.0,
            parser_name=self._parser_name,
            parser_version=self._parser_version
        )

        return result

    def _parse_content_impl(
        self,
        content: Union[str, bytes],
        document_id: str,
        document_type: DocumentType,
        **metadata_kwargs
    ) -> ParsedDocument:
        """
        Parse DOCX content directly.

        Args:
            content: DOCX content as bytes
            document_id: Document identifier
            document_type: Type of document
            **metadata_kwargs: Additional metadata fields

        Returns:
            ParsedDocument: Parsed document result
        """
        if isinstance(content, str):
            raise ValueError("DOCX content must be provided as bytes")

        # Create temporary file for parsing
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_file.write(content)
            tmp_path = Path(tmp_file.name)

        try:
            # Parse the temporary file
            result = self._parse_file_impl(tmp_path)

            # Update metadata with provided information
            result.metadata.document_id = document_id
            for key, value in metadata_kwargs.items():
                if hasattr(result.metadata, key):
                    setattr(result.metadata, key, value)
                else:
                    result.metadata.add_custom_field(key, value)

            return result
        finally:
            # Clean up temporary file
            try:
                tmp_path.unlink()
            except Exception:
                pass

    def _extract_docx_properties(self, doc: Document, metadata: DocumentMetadata) -> None:
        """Extract properties from DOCX document."""
        try:
            core_props = doc.core_properties

            # Extract basic properties
            if core_props.title:
                metadata.title = core_props.title

            if core_props.author:
                metadata.author = core_props.author

            if core_props.subject:
                metadata.subject = core_props.subject

            if core_props.keywords:
                # Split keywords by common separators
                keywords = [k.strip() for k in core_props.keywords.replace(',', ';').split(';') if k.strip()]
                metadata.keywords = keywords

            # Extract dates
            if core_props.created:
                metadata.created_at = core_props.created

            if core_props.modified:
                metadata.modified_at = core_props.modified

            # Add additional properties as custom fields
            if core_props.category:
                metadata.add_custom_field('category', core_props.category)

            if core_props.comments:
                metadata.add_custom_field('comments', core_props.comments)

            if core_props.content_status:
                metadata.add_custom_field('content_status', core_props.content_status)

            if core_props.identifier:
                metadata.add_custom_field('identifier', core_props.identifier)

            if core_props.language:
                metadata.language = core_props.language

            if core_props.last_modified_by:
                metadata.add_custom_field('last_modified_by', core_props.last_modified_by)

            if core_props.last_printed:
                metadata.add_custom_field('last_printed', core_props.last_printed)

            if core_props.revision:
                metadata.add_custom_field('revision', core_props.revision)

            if core_props.version:
                metadata.add_custom_field('version', core_props.version)

        except Exception as e:
            self.logger.warning(f"Failed to extract DOCX properties: {e}")

    def _extract_hyperlinks(self, doc: Document) -> List[Dict[str, str]]:
        """Extract hyperlinks from DOCX document."""
        links = []

        try:
            # Get hyperlink relationships
            for rel in doc.part.rels.values():
                if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink":
                    links.append({
                        'url': rel.target_ref,
                        'type': 'hyperlink',
                        'relationship_id': rel.rId
                    })
        except Exception as e:
            self.logger.warning(f"Failed to extract hyperlinks: {e}")

        return links
